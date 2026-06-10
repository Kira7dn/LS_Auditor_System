import sys
from pathlib import Path
import docx

def convert_docx_to_md(docx_path: Path, md_path: Path) -> None:
    """Converts a DOCX file to Markdown format using python-docx."""
    doc = docx.Document(docx_path)
    md_content = []
    
    # Process paragraphs
    for p in doc.paragraphs:
        text = p.text
        if not text.strip():
            md_content.append("")
            continue
            
        style = p.style.name if p.style else ""
        
        # Check formatting of runs to preserve bold/italic if possible
        formatted_runs = []
        for run in p.runs:
            run_text = run.text
            if run.bold:
                run_text = f"**{run_text}**"
            if run.italic:
                run_text = f"*{run_text}*"
            formatted_runs.append(run_text)
        text = "".join(formatted_runs)
        
        # Determine heading level
        if style.startswith("Heading 1"):
            md_content.append(f"\n# {text}\n")
        elif style.startswith("Heading 2"):
            md_content.append(f"\n## {text}\n")
        elif style.startswith("Heading 3"):
            md_content.append(f"\n### {text}\n")
        elif style.startswith("Heading 4"):
            md_content.append(f"\n#### {text}\n")
        elif "bullet" in style.lower():
            md_content.append(f"* {text}")
        elif "number" in style.lower():
            md_content.append(f"1. {text}")
        else:
            md_content.append(text)
            
    # Process tables if any
    for i, table in enumerate(doc.tables):
        md_content.append(f"\n\n### Bảng trích xuất {i+1}\n\n")
        for r_idx, row in enumerate(table.rows):
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            md_content.append("| " + " | ".join(cells) + " |")
            if r_idx == 0:
                md_content.append("| " + " | ".join(["---"] * len(cells)) + " |")
        md_content.append("\n")
        
    with open(md_path, "w", encoding="utf-8") as f:
        f.write("\n".join(md_content))
    print(f"Converted {docx_path} to {md_path}")

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python convert_docx_to_md.py <input_docx_path> <output_md_path>")
        sys.exit(1)
        
    convert_docx_to_md(Path(sys.argv[1]), Path(sys.argv[2]))
